import os
import platform
import subprocess
from docx import Document
from docx.shared import Pt


def fetchPrevCompany() -> str:
    with open('prevCompany.txt', 'r') as previous:
        x = previous.read()

        if len(x.strip()) < 1:
            newPrevCompany: str = input('Enter current company on cover letter: ')
            changePrevCompany(newPrevCompany.strip())
            return newPrevCompany.strip()

        return x.strip()

def get_soffice_path() -> str:
    system = platform.system()

    if system == 'Darwin':  # macOS
        return '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    
    elif system == 'Windows': 
        return r'C:\Program Files\LibreOffice\program\soffice.exe'
    elif system == 'Linux':
        return '/usr/bin/soffice'
    
    else:
        raise Exception(f'Unsupported operating system: {system}')

def changePrevCompany(newCompany: str) -> None:
    with open('prevCompany.txt', 'w') as previous:
        previous.write(newCompany)

def fetchOutputDir() -> str:
    with open('outputPath.txt', 'r') as outputDir:
        return outputDir.read()
    
def fetchInputDir() -> str:
    with open('inputPath.txt', 'r') as inputDir:
        return inputDir.read()

doc = Document(fetchInputDir())

font_name = "Calibri"
search_text = fetchPrevCompany()
replace_text = input("Enter new company name: ")
changePrevCompany(replace_text)

for paragraph in doc.paragraphs:
    if search_text in paragraph.text:
        paragraph.clear()
        run = paragraph.add_run(replace_text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = font_name
        break


current_dir = os.path.abspath('')
output_dir = os.path.normpath(os.path.join(current_dir, fetchOutputDir()))
output_pdf_path = os.path.join(fetchOutputDir())

# Save the changes in the original docx file (if preferred)
doc.save(fetchInputDir())


# Convert DOCX to PDF
command = [
    get_soffice_path(),
    '--headless',
    '--convert-to',
    'pdf',
    '--outdir',
    output_pdf_path,
    fetchInputDir()
]

subprocess.run(command, check=True)